"""
resume_loader.py — Extract plain text from a .docx resume file.

Requires: pip install python-docx
"""

import os


def load_resume(filepath: str = "resume.docx") -> str:
    """
    Load and extract all text from a .docx file.

    Args:
        filepath: Path to the Word document (default: resume.docx).

    Returns:
        A single string containing the full resume text.

    Raises:
        FileNotFoundError: If the file does not exist.
        ImportError:       If python-docx is not installed.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(
            f"Resume file not found: '{filepath}'\n"
            "Please place your resume.docx in the project root directory."
        )

    if not filepath.endswith(".docx"):
        raise ValueError(
            f"Expected a .docx file, got: '{filepath}'\n"
            "Please convert your resume to Word (.docx) format first."
        )

    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    doc = Document(filepath)

    sections = []

    # ── Paragraphs (body text, headings, bullets) ────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # ── Tables (skills grid, education table, etc.) ──────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                sections.append(row_text)

    if not sections:
        raise ValueError(
            f"No text could be extracted from '{filepath}'. "
            "Ensure the document contains readable text (not just images)."
        )

    full_text = "\n".join(sections)
    print(f"[resume_loader] Loaded resume: {len(full_text)} characters, "
          f"{len(sections)} paragraphs.")
    return full_text
